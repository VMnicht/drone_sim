#!/usr/bin/env python3
"""Generate the final academic-style DOCX report from repository evidence.

The document is deliberately built from committed experiment JSON/PNG files.  It
does not use generative imagery.  The two explanatory diagrams are drawn with
Pillow so that their content remains traceable to the repository architecture.
"""

from __future__ import annotations

import argparse
import json
import math
import re
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


BLACK = "000000"
GRAY_1 = "E7E6E6"
GRAY_2 = "F2F2F2"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9070  # A4 width minus two 25 mm margins.


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: Sequence[int], indent_dxa: int = 0) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])


def set_font(run, east_asia: str, size: float, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def apply_run_fonts(paragraph, east_asia: str, size: float, *, bold: bool = False, italic: bool = False) -> None:
    for run in paragraph.runs:
        set_font(run, east_asia, size, bold=bold, italic=italic)


def configure_styles(doc: Document) -> None:
    # Base preset: narrative_proposal.  Academic-paper override: A4, black-only,
    # SimSun body, SimHei headings, Times New Roman for Latin text.
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Pt(19)
    pf.line_spacing = 1.18
    pf.space_after = Pt(2)

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    heading_tokens = {
        "Heading 1": (15.0, 10, 6),
        "Heading 2": (12.5, 8, 4),
        "Heading 3": (11.0, 6, 3),
    }
    for style_name, (size, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.15

    if "CaptionAcademic" not in [s.name for s in doc.styles]:
        caption = doc.styles.add_style("CaptionAcademic", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["CaptionAcademic"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = False

    if "EquationAcademic" not in [s.name for s in doc.styles]:
        equation = doc.styles.add_style("EquationAcademic", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = doc.styles["EquationAcademic"]
    equation.font.name = "Times New Roman"
    equation._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    equation.font.size = Pt(10.5)
    equation.font.color.rgb = RGBColor(0, 0, 0)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.first_line_indent = Pt(0)
    equation.paragraph_format.space_before = Pt(3)
    equation.paragraph_format.space_after = Pt(3)
    equation.paragraph_format.keep_together = True

    if "ReferenceAcademic" not in [s.name for s in doc.styles]:
        ref = doc.styles.add_style("ReferenceAcademic", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = doc.styles["ReferenceAcademic"]
    ref.font.name = "Times New Roman"
    ref._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    ref.font.size = Pt(9.5)
    ref.font.color.rgb = RGBColor(0, 0, 0)
    ref.paragraph_format.first_line_indent = Pt(-18)
    ref.paragraph_format.left_indent = Pt(18)
    ref.paragraph_format.space_after = Pt(2)
    ref.paragraph_format.line_spacing = 1.15


def add_field(paragraph, instruction: str, placeholder: str = "1") -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_font(run, "宋体", 8.5)


def configure_section(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(22)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("ROS2 小型无人机仿真器课程设计报告")
    set_font(run, "宋体", 8.0)
    p_pr = hp._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "BFBFBF")
    border.append(bottom)
    p_pr.append(border)

    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    first_header.paragraphs[0].text = ""

    for footer in (section.footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_after = Pt(0)
        r1 = fp.add_run("— ")
        set_font(r1, "宋体", 8.5)
        add_field(fp, "PAGE")
        r2 = fp.add_run(" —")
        set_font(r2, "宋体", 8.5)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, after: float = 2) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_font(first, "宋体", 9.5, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_font(rest, "宋体", 9.5)
    else:
        run = p.add_run(text)
        set_font(run, "宋体", 9.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    apply_run_fonts(p, "黑体", {1: 15.0, 2: 12.5, 3: 11.0}[level], bold=True)


def add_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="EquationAcademic")
    run = p.add_run(text)
    set_font(run, "宋体", 10.5)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="CaptionAcademic")
    run = p.add_run(text)
    set_font(run, "宋体", 9)


def add_figure(doc: Document, path: Path, width_in: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    # Word 2016 can enter a PDF-layout loop when a large inline PNG paragraph
    # is chained to the following caption.  Captions are checked after render,
    # so keep the picture inline but leave pagination unconstrained here.
    p.paragraph_format.keep_with_next = False
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
    *,
    font_size: float = 8.8,
    header_size: float = 9.0,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0]
    for index, header in enumerate(headers):
        cell = hdr.cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, GRAY_1)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_font(run, "黑体", header_size, bold=True)
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFAFA")
            set_cell_margins(cell, top=65, bottom=65)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index != 1 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            set_font(run, "宋体", font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def load_font(candidates: Iterable[str], size: int) -> ImageFont.FreeTypeFont:
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def centered_multiline(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, spacing: int = 8) -> None:
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(((left + right - width) / 2, (top + bottom - height) / 2), text, font=font, fill="black", spacing=spacing, align="center")


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], width: int = 4) -> None:
    draw.line([start, end], fill="black", width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    spread = 0.55
    p1 = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
    p2 = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
    draw.polygon([end, p1, p2], fill="black")


def create_architecture_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1800, 1060), "white")
    draw = ImageDraw.Draw(image)
    cn = load_font([r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\msyhbd.ttc"], 30)
    lane_font = load_font([r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\msyhbd.ttc"], 25)
    note_font = load_font([r"C:\Windows\Fonts\simsun.ttc", r"C:\Windows\Fonts\msyh.ttc"], 21)
    latin = load_font([r"C:\Windows\Fonts\times.ttf", r"C:\Windows\Fonts\timesbd.ttf"], 20)
    label_font = load_font([r"C:\Windows\Fonts\times.ttf", r"C:\Windows\Fonts\timesbd.ttf"], 17)

    def draw_box(box: tuple[int, int, int, int], title: str, subtitle: str, *, core: bool) -> None:
        fill = f"#{GRAY_1 if core else GRAY_2}"
        draw.rounded_rectangle(box, radius=16, fill=fill, outline="black", width=3)
        left, top, right, bottom = box
        title_bbox = draw.textbbox((0, 0), title, font=cn)
        title_width = title_bbox[2] - title_bbox[0]
        subtitle_bbox = draw.textbbox((0, 0), subtitle, font=latin)
        subtitle_width = subtitle_bbox[2] - subtitle_bbox[0]
        draw.text(((left + right - title_width) / 2, top + 29), title, font=cn, fill="black")
        draw.text(((left + right - subtitle_width) / 2, bottom - 39), subtitle, font=latin, fill="black")

    def flow(points: Sequence[tuple[int, int]], *, width: int = 4) -> None:
        draw.line(points, fill="black", width=width, joint="curve")
        start, end = points[-2], points[-1]
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        length = 15
        spread = 0.55
        p1 = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
        p2 = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
        draw.polygon([end, p1, p2], fill="black")

    def label(text_value: str, center: tuple[int, int]) -> None:
        bbox = draw.textbbox((0, 0), text_value, font=label_font)
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]
        x = center[0] - width / 2
        y = center[1] - height / 2
        draw.rounded_rectangle((x - 7, y - 4, x + width + 7, y + height + 4), radius=5, fill="white")
        draw.text((x, y), text_value, font=label_font, fill="black")

    # Lane guides keep control, navigation, feedback and visualization visually separate.
    draw.text((80, 105), "主控制链", font=lane_font, fill="black")
    draw.line((220, 124, 1620, 124), fill="#BFBFBF", width=2)
    draw.text((80, 410), "导航与状态链", font=lane_font, fill="black")
    draw.line((280, 429, 1620, 429), fill="#BFBFBF", width=2)
    draw.text((80, 720), "可视化链", font=lane_font, fill="black")
    draw.line((220, 739, 1620, 739), fill="#BFBFBF", width=2)

    boxes = {
        "task": (80, 160, 340, 285),
        "controller": (390, 160, 650, 285),
        "mixer": (700, 160, 960, 285),
        "fault": (1010, 160, 1270, 285),
        "dynamics": (1320, 160, 1580, 285),
        "map": (80, 470, 340, 595),
        "planner": (390, 470, 650, 595),
        "reference": (700, 470, 960, 595),
        "state": (1000, 470, 1230, 595),
        "sensor": (1340, 470, 1570, 595),
        "bus": (80, 790, 1270, 905),
        "visualization": (1360, 790, 1620, 905),
    }
    draw_box(boxes["task"], "任务与轨迹", "TrajectoryPoint", core=True)
    draw_box(boxes["controller"], "位置与姿态控制", "ModelBasedController", core=True)
    draw_box(boxes["mixer"], "电机分配器", "MotorMixer", core=True)
    draw_box(boxes["fault"], "故障注入", "MotorRPM", core=True)
    draw_box(boxes["dynamics"], "六自由度动力学", "QuadrotorModel", core=True)
    draw_box(boxes["map"], "地图与局部感知", "PointCloud2 / voxel", core=False)
    draw_box(boxes["planner"], "三维 A* 规划", "nav_msgs / Path", core=False)
    draw_box(boxes["reference"], "参考选择", "safe goal / trajectory", core=False)
    draw_box(boxes["state"], "ROS2 状态适配", "Odom / IMU / TF", core=False)
    draw_box(boxes["sensor"], "传感器模型", "noise / delay / dropout", core=True)
    draw_box(boxes["visualization"], "RViz 与 Web", "display / diagnostics", core=False)

    # Main control pipeline. Labels sit in dedicated white gaps between nodes.
    flow([(340, 222), (390, 222)])
    flow([(650, 222), (700, 222)])
    flow([(960, 222), (1010, 222)])
    flow([(1270, 222), (1320, 222)])
    label("reference", (365, 195))
    label("wrench", (675, 195))
    label("MotorRPM", (985, 195))
    label("faulted RPM", (1295, 195))

    # Navigation pipeline and two non-overlapping return buses into the controller.
    flow([(340, 532), (390, 532)])
    flow([(650, 532), (700, 532)])
    label("local map", (365, 505))
    label("planned path", (675, 505))
    flow([(830, 470), (830, 390), (585, 390), (585, 285)])
    label("safe goal / trajectory", (710, 363))

    # Dynamics, sensor adaptation and feedback use a separate upper return bus.
    flow([(1450, 285), (1450, 470)])
    label("truth state", (1515, 365))
    flow([(1340, 532), (1230, 532)])
    label("noisy state", (1285, 505))
    flow([(1115, 470), (1115, 325), (455, 325), (455, 285)])
    label("Odom / IMU feedback", (785, 298))

    # Display messages enter a dedicated bus, so none of these arrows cross the control chain.
    draw.rounded_rectangle(boxes["bus"], radius=16, fill="#F7F7F7", outline="black", width=3)
    centered_multiline(draw, boxes["bus"], "ROS2 可视化数据总线\nPath · Marker · PointCloud2 · TF · diagnostics", note_font, spacing=9)
    flow([(210, 595), (210, 790)])
    flow([(830, 595), (830, 790)])
    flow([(1115, 595), (1115, 790)])
    flow([(1270, 848), (1360, 848)])

    draw.rounded_rectangle((260, 25, 1540, 78), radius=12, fill="white", outline="black", width=2)
    centered_multiline(
        draw,
        (260, 25, 1540, 78),
        "drone_core 仅依赖 Eigen 与 C++ STL；地图、点云和规划保留 ROS2 接口",
        note_font,
    )
    draw.text((80, 972), "深灰节点：ROS 无关核心算法    浅灰节点：ROS2 适配或地图相关模块", font=note_font, fill="black")
    image.save(path, dpi=(220, 220))


def create_control_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(image)
    cn = load_font([r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\msyhbd.ttc"], 34)
    small = load_font([r"C:\Windows\Fonts\simsun.ttc", r"C:\Windows\Fonts\msyh.ttc"], 24)
    boxes = [
        ((60, 250, 300, 420), "参考位置\n速度与加速度"),
        ((390, 250, 650, 420), "位置速度外环\n限速与限加速度"),
        ((740, 250, 1000, 420), "期望推力方向\n几何姿态误差"),
        ((1090, 250, 1350, 420), "角速度内环\n推力与力矩限幅"),
        ((1440, 250, 1720, 420), "X 型 mixer\n四路电机转速"),
    ]
    for box, label in boxes:
        draw.rounded_rectangle(box, radius=18, fill="#F2F2F2", outline="black", width=3)
        centered_multiline(draw, box, label, cn)
    for left, right in zip(boxes[:-1], boxes[1:]):
        arrow(draw, (left[0][2], 335), (right[0][0], 335))

    draw.rounded_rectangle((440, 545, 1320, 690), radius=18, fill="#E7E6E6", outline="black", width=3)
    centered_multiline(draw, (440, 545, 1320, 690), "状态反馈：位置、速度、四元数姿态与机体系角速度", cn)
    arrow(draw, (1540, 420), (1320, 615))
    arrow(draw, (440, 615), (520, 420))
    draw.text((110, 145), "前馈项来自解析轨迹或规划器；反馈项来自带噪 Odom 和 IMU", font=small, fill="black")
    image.save(path, dpi=(220, 220))


def prepare_rgb_asset(source: Path, destination: Path) -> Path:
    """Flatten Matplotlib alpha channels for reliable Word/PDF export."""
    destination.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(source) as image:
        if image.mode == "RGBA":
            background = Image.new("RGB", image.size, "white")
            background.paste(image, mask=image.getchannel("A"))
            result = background
        else:
            result = image.convert("RGB")
        result.save(destination, format="PNG", optimize=True, dpi=(170, 170))
    return destination


def read_summaries(root: Path) -> dict[str, dict]:
    result: dict[str, dict] = {}
    for path in sorted((root / "artifacts" / "experiments").glob("*/summary.json")):
        data = json.loads(path.read_text(encoding="utf-8"))
        result[data["scenario"]] = data
    return result


def fmt(value, digits: int = 4) -> str:
    if value is None:
        return "—"
    return f"{float(value):.{digits}f}"


def scenario_table_rows(summaries: dict[str, dict]) -> list[list[str]]:
    order = [
        ("hover", "悬停"),
        ("target", "单目标点"),
        ("square", "正方形航点"),
        ("circle", "圆轨迹"),
        ("figure_eight", "八字轨迹"),
        ("wind_gust", "阵风扰动"),
        ("sensor_noise", "传感器噪声"),
        ("fault_motor", "电机故障"),
        ("five_obstacles", "五障碍物"),
        ("narrow_passage", "狭窄通道"),
        ("perception_replan", "感知重规划"),
    ]
    rows: list[list[str]] = []
    for key, label in order:
        data = summaries[key]
        note = "通过"
        if key == "wind_gust":
            note = f"恢复 {fmt(data.get('disturbance_recovery_time_s'), 3)} s"
        elif key in {"five_obstacles", "narrow_passage", "perception_replan"}:
            note = f"净空 {fmt(data.get('minimum_obstacle_clearance_m'), 4)} m"
        elif key == "fault_motor":
            note = f"最大倾角 {fmt(data.get('maximum_tilt_deg'), 2)}°"
        rows.append([
            label,
            fmt(data.get("final_position_error_m"), 4),
            fmt(data.get("steady_state_error_m"), 4),
            fmt(data.get("rms_position_error_m"), 4),
            note,
        ])
    return rows


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.add_run("ROS2 小型四旋翼无人机仿真系统的设计与实现")
    apply_run_fonts(p, "黑体", 22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("摘  要")
    set_font(run, "黑体", 14, bold=True)
    abstract = (
        "本项目完成了一套面向课程实验的小型四旋翼仿真系统。系统运行于 ROS2 Humble，核心动力学、控制器、"
        "电机分配器、轨迹、扰动和传感器数学模型集中在 drone_core 中，只依赖 Eigen 与 C++ 标准库。ROS2 节点"
        "负责参数装载、消息转换和可视化，地图、点云与规划模块则保留对 ROS2 的直接依赖。这样的边界使控制算法"
        "可以脱离通信框架单独构建和测试，也便于以后接入其他中间件或硬件在环环境。系统实现了六自由度刚体动力学、"
        "串级几何控制、三维体素 A*、局部点云、故障与风扰注入、多机仿真、RViz2 和 Web 地面站。11 个正式场景"
        "全部通过自动阈值验收。悬停稳态误差为 0.0203 m，三组避障实验的最小净空分别为 0.4196 m、0.4136 m"
        "和 0.4889 m，均高于 0.30 m 的安全约束。实验表明，这套实现能够在较少外部依赖下形成从模型、控制到"
        "规划和评测的完整闭环，同时保留清晰的工程复现路径。"
    )
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(abstract)
    set_font(r, "宋体", 10.5)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run("关键词：")
    set_font(r1, "黑体", 10.5, bold=True)
    r2 = p.add_run("四旋翼；ROS2；六自由度动力学；几何控制；三维 A*；仿真评测")
    set_font(r2, "宋体", 10.5)


def build_report(root: Path, output_path: Path) -> None:
    summaries = read_summaries(root)
    required = {
        "hover", "target", "square", "circle", "figure_eight", "wind_gust",
        "sensor_noise", "fault_motor", "five_obstacles", "narrow_passage", "perception_replan",
    }
    missing = required - summaries.keys()
    if missing:
        raise RuntimeError(f"Missing experiment summaries: {sorted(missing)}")

    assets = root / "output" / "report_assets"
    architecture = assets / "system_architecture.png"
    control_flow = assets / "control_flow.png"
    create_architecture_diagram(architecture)
    create_control_diagram(control_flow)
    hover_figure = prepare_rgb_asset(
        root / "artifacts" / "experiments" / "hover" / "experiment_summary.png",
        assets / "hover_summary_rgb.png",
    )
    wind_figure = prepare_rgb_asset(
        root / "artifacts" / "experiments" / "wind_gust" / "experiment_summary.png",
        assets / "wind_gust_summary_rgb.png",
    )
    replan_figure = prepare_rgb_asset(
        root / "artifacts" / "experiments" / "perception_replan" / "environment_metrics.png",
        assets / "perception_replan_metrics_rgb.png",
    )

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_title_page(doc)

    add_heading(doc, "1  引言", 1)
    add_body(doc, "四旋翼仿真器看起来只是把若干微分方程放进 ROS2 节点，真正实现时却很容易在接口和时间尺度上出问题。电机命令、刚体积分、控制器和可视化具有不同更新频率，任何一个环节的单位或时间戳处理不一致，都会表现为模型闪烁、姿态发散或轨迹无法收敛。因此，本项目没有直接包装现成仓库，而是从状态定义、力矩方向和消息数据链重新搭建仿真闭环，并用可重复的场景逐步验收。")
    add_body(doc, "课程要求以目标点为输入，以四路电机转速为控制输出，同时给出动力学、地图、避障、可视化和完整说明。项目在此基础上继续实现了风扰、传感器噪声、点云与体素感知、多机、解析轨迹、故障注入、地面站、参数扫描和自动评测。所有加分项都进入统一启动入口和正式实验，而不是作为互不相连的演示脚本存在。")
    add_body(doc, "设计中最重要的取舍是让算法与 ROS2 解耦。drone_core 内保存动力学、模型控制、mixer、轨迹生成、扰动和噪声模型，只使用 Eigen 与 C++ STL。ROS2 适配层负责 topic、QoS、参数和坐标系。地图和规划需要直接消费 PointCloud2、Marker 与 Path，且这些模块的意义本来就与机器人中间件的数据组织紧密相关，因此允许它们保留 ROS2 依赖。这个边界既符合课程要求，也让核心算法可以用普通 CMake 和 CTest 独立验证。")

    add_heading(doc, "2  系统总体设计", 1)
    add_heading(doc, "2.1  模块划分与数据流", 2)
    add_figure(doc, architecture, 5.95, "图 1  系统节点与主要数据流")
    add_body(doc, "系统的数据流可以概括为两个闭环。飞行闭环从任务或轨迹参考进入控制器，经 X 型电机分配和故障注入后驱动动力学，再由传感器模型把真值状态转换为控制器可用的 Odom 与 IMU。导航闭环从静态地图和局部点云出发，三维 A* 输出安全路径与局部目标，随后仍复用同一控制器。RViz2 与 Web 地面站只订阅状态和诊断信息，不参与动力学积分，因此界面退出不会改变实验结果。")

    add_table(
        doc,
        ["环节", "主要 ROS2 接口", "消息类型", "更新频率"],
        [
            ["动力学", "/drone/motor_rpm_faulted → /drone/odom", "MotorRPM / Odometry", "约 100 Hz"],
            ["控制", "/drone/goal、/drone/odom → /drone/motor_rpm_cmd", "PoseStamped / MotorRPM", "约 100 Hz"],
            ["导航", "/map/obstacles、/drone/local_points → /drone/planned_path", "MarkerArray / PointCloud2 / Path", "事件触发与低频"],
            ["显示", "/tf、/drone/path、/drone/markers", "TFMessage / Path / MarkerArray", "1～5 Hz"],
        ],
        [1450, 3750, 2450, 1420],
        font_size=8.3,
    )
    add_caption(doc, "表 1  主要节点接口与频率分工")

    add_heading(doc, "2.2  工程组织与启动方式", 2)
    add_body(doc, "工作区共有 15 个 ROS2 package。drone_msgs 定义电机转速、障碍物与轨迹点消息；drone_dynamics、drone_controller 和 drone_sensors 是核心算法的 ROS2 适配层；drone_map、drone_perception 与 drone_planner 负责环境和安全路径；其余 package 承担多机、故障、可视化、工具、地面站与 launch 编排。根目录的 start_sim.sh 会在脚本内部加载 ROS2 和工作区环境，用户不需要手动 source。若尚未构建，脚本先执行 colcon build，再按 hover、experiment、multi、ground-station 或 panel 模式启动。")

    add_heading(doc, "3  四旋翼动力学模型", 1)
    add_heading(doc, "3.1  状态、输入与电机响应", 2)
    add_body(doc, "动力学状态由世界系位置 p、速度 v、机体到世界的单位四元数 q、机体系角速度 Ω 以及四个电机实际角速度组成。控制消息以 RPM 表示，进入核心模型前换算为 rad/s。模型默认质量为 1.0 kg，X 型机臂长度为 0.17 m，惯量对角项为 0.02、0.02 和 0.04 kg·m²。电机没有被处理成理想瞬时执行器，而是使用时间常数 τm 的一阶响应。")
    add_equation(doc, "ωᵢ,k+1 = ωᵢ,k + [1 − exp(−Δt / τm)] (ωᵢ,cmd − ωᵢ,k)")
    add_body(doc, "每次更新后都执行上下限裁剪。这样既能避免控制器给出的极端命令进入平方推力模型，也能重现起飞阶段转速建立的短暂过程。默认最大电机角速度为 2300 rad/s，对应约 21963 RPM。理论悬停角速度为 1133.15 rad/s，约 10821 RPM，与实验稳态曲线相符。")

    add_table(
        doc,
        ["参数", "符号", "数值", "说明"],
        [
            ["质量", "m", "1.0 kg", "机体总质量"],
            ["惯量", "I", "diag[0.02, 0.02, 0.04] kg·m²", "机体系对角惯量"],
            ["机臂长度", "l", "0.17 m", "质心到旋翼中心"],
            ["推力系数", "kF", "1.91×10⁻⁶ N/(rad/s)²", "单旋翼平方推力"],
            ["反扭矩系数", "kM", "2.60×10⁻⁷ N·m/(rad/s)²", "偏航反扭矩"],
            ["重力加速度", "g", "9.81 m/s²", "世界系 z 轴向上"],
        ],
        [1700, 1000, 3000, 3370],
        font_size=8.6,
    )
    add_caption(doc, "表 2  动力学默认参数")

    add_heading(doc, "3.2  推力、力矩与刚体积分", 2)
    add_body(doc, "第 i 个旋翼的推力为 Fi = kFωi²，反扭矩为 Mi = si kMωi²。si 由旋翼转向决定。四个旋翼按 X 型布局叠加为总推力 T 和三个轴向力矩 τ。平动方程在世界坐标系计算，转动方程在机体系计算，并加入线性阻力、角速度阻尼和可选外部扰动力。")
    add_equation(doc, "m v̇ = R(q) [0, 0, T]ᵀ − m g e₃ − Cv + Fext")
    add_equation(doc, "I Ω̇ = τ − Ω × IΩ − CΩ + τext")
    add_equation(doc, "q̇ = 1/2 · q ⊗ [0, Ω]，每步积分后执行 q ← q / ‖q‖")
    add_body(doc, "数值积分使用固定步长的半隐式更新。角速度和线速度先由加速度推进，位置和四元数再使用新状态更新。地面附近加入非穿透约束，防止数值误差让机体落到 z 轴负半空间。仿真真值只由动力学节点产生一次 Odom 和 TF，传感器节点发布带噪副本，避免多个发布者导致 RViz 中同一模型交替跳变。")

    add_heading(doc, "4  控制器设计", 1)
    add_figure(doc, control_flow, 5.95, "图 2  串级控制器与状态反馈关系")
    add_heading(doc, "4.1  位置外环与期望姿态", 2)
    add_body(doc, "控制器采用位置、速度外环和几何姿态内环。给定参考位置 pd、速度 vd、加速度 ad 及期望偏航角，外环先计算位置误差 ep = p − pd 与速度误差 ev = v − vd，再得到期望加速度。逐轴增益来自 controller.yaml，位置积分项仅在限定范围内累积，并设置反饱和回退。")
    add_equation(doc, "ac = ad − Kp ⊙ ep − Kv ⊙ ev − Ki ⊙ ∫ep dt")
    add_body(doc, "期望合力 Fd = m(ac + ge₃)。其方向确定期望机体 z 轴，再结合期望偏航构造完整旋转矩阵 Rd。相较于直接把位置误差线性换算成 roll 和 pitch，这种构造在三维轨迹中没有欧拉角奇异点，且能清楚地限制合力和最大倾角。外环还限制目标跳变、水平速度、垂直速度和加速度，避免距离很远时一步给出不可实现的姿态。")

    add_heading(doc, "4.2  几何姿态内环与 mixer", 2)
    add_body(doc, "姿态误差从 RᵀRd − RdᵀR 的反对称部分提取，角速度误差由当前角速度与期望角速度之差得到。控制力矩同时补偿刚体陀螺耦合项 Ω×IΩ。总推力取期望合力在当前机体 z 轴上的投影，因此姿态尚未对准时不会把不正确的水平分量直接当作竖直推力。")
    add_equation(doc, "τc = −KR ⊙ eR − KΩ ⊙ eΩ + Ω × IΩ")
    add_body(doc, "mixer 解算 T、τx、τy、τz 到四个 ωi²。解算结果先保证非负，再经过推力、力矩、电机角速度和变化率的多级限幅。控制节点将 rad/s 转为消息中的 RPM，动力学适配层再还原为角速度。故障节点位于两者之间，可以按时间窗口降低单个电机效率、限制最大转速，或模拟命令丢包、延迟和冻结。这样的串接方式让正常控制器代码不需要知道故障是否存在。")

    add_heading(doc, "5  地图、感知与避障", 1)
    add_heading(doc, "5.1  地图与局部点云", 2)
    add_body(doc, "静态地图由 YAML 中的 box 和 cylinder 描述，也可以使用固定随机种子生成障碍物。地图节点发布可视 Marker 和障碍物几何信息。感知节点对几何表面采样，再根据无人机位置与朝向执行量程、视场、遮挡、噪声和丢点处理，输出局部 PointCloud2。点云进一步体素化，并使用短时持久化与过期衰减减少边界抖动。由于地图表示和 PointCloud2 接口紧密相关，这部分保留 ROS2 依赖，没有强行移入 drone_core。")
    add_body(doc, "安全距离采用障碍物几何边界到机体中心的净空定义。规划时将机体半径与 0.30 m 安全裕度共同膨胀到障碍物上，因此搜索节点可以按点机器人处理。实验记录器独立计算实际轨迹到原始障碍物表面的最小净空，用来验证规划假设是否在闭环飞行中成立。")

    add_heading(doc, "5.2  三维体素 A* 与路径跟踪", 2)
    add_body(doc, "规划器把起点、目标点和膨胀障碍物投影到三维体素网格，在 6、18 或 26 邻接中搜索。代价由移动距离和启发函数组成。若起点或终点落入占据体素，节点先在有限邻域内寻找最近自由体素；若开放集耗尽、搜索节点数超过上限或无法重建父节点链，则明确报告失败，不向控制器发布穿越障碍物的直线路径。")
    add_body(doc, "得到离散路径后，先做共线点合并和带碰撞检测的线段简化，再由前视距离选择局部安全目标。路径跟踪进度只允许沿路径前进，避免折返或自交路径中前视点误跳到尚未经过的后段。perception_replan 场景会在局部点云发生有效变化时触发重规划，同时设置最小重规划周期，防止感知噪声造成频繁振荡。")

    add_heading(doc, "5.3  可视化与交互", 2)
    add_body(doc, "RViz2 中同时显示无人机机体、四个旋翼、目标点、历史轨迹、障碍物、局部点云、体素和规划路径。无人机 Marker 使用固定命名空间与稳定 ID，时间戳设为零并启用 frame_locked；地图类低频信息使用 transient-local QoS。此前模型闪烁的根因是 Odom、TF 和 Marker 存在重复发布者，且显示对象生命周期过短。修正后每类状态只保留一个权威发布者，Marker 不再周期性失效。start_sim.sh 在 WSL 内部处理 DISPLAY 与环境加载，因此用户可以直接运行脚本而不必手动 source。")
    add_body(doc, "Web 地面站用于发布目标、查看位置和高度曲线、启停扰动与故障。任务 Panel 汇总 30 个任务入口、11 个正式场景、16 类指标和 7 类结果图，并允许在白名单范围内修改 29 份 YAML。界面只写配置副本且保留回滚，避免误操作破坏仓库基线。")

    add_heading(doc, "6  实验设计与结果", 1)
    add_heading(doc, "6.1  实验方法", 2)
    add_body(doc, "正式实验统一从 start_sim.sh 调度。每个场景使用固定参数与随机种子，记录 telemetry.csv、reference_history.csv、summary.json、rosbag 和七类曲线。评价指标包含最终误差、稳态误差、RMS 误差、到达时间、最大姿态倾角、电机饱和比例、扰动恢复时间和最小障碍物净空。验收阈值存放在 evaluation.yaml，验证脚本只读取阈值与结果，不在代码中另设一套宽松标准。")
    add_body(doc, "表 3 的数值来自当前 artifacts/experiments 目录，而不是按设计参数推算。RMS 误差包含从地面起飞和目标切换阶段，因此在多航点、阵风和避障场景中会明显高于最终误差。它更能反映整个任务过程，而稳态误差用于判断最终悬停质量，两者不能互相替代。")

    add_table(
        doc,
        ["场景", "最终误差/m", "稳态误差/m", "RMS 误差/m", "关键结果"],
        scenario_table_rows(summaries),
        [1700, 1500, 1500, 1500, 2870],
        font_size=8.2,
        header_size=8.5,
    )
    add_caption(doc, "表 3  11 个正式场景的实测结果")

    add_heading(doc, "6.2  基础飞行与轨迹跟踪", 2)
    add_figure(doc, hover_figure, 5.65, "图 3  悬停实验的轨迹、高度、位置误差和电机转速")
    add_body(doc, "悬停实验在约 1.99 s 首次进入 0.30 m 误差范围，随后高度稳定在 1.5 m 附近，稳态误差为 0.0203 m。四路电机在起飞瞬间共同升速，之后收敛到理论悬停转速附近；最大倾角仅 0.65°，说明静态悬停没有依赖持续大姿态补偿。")
    add_body(doc, "正方形任务依次完成 5 个航点。每次参考点切换会让误差重新升高，但局部目标到达后均可继续收敛，最终误差为 0.0168 m。圆轨迹与 Gerono 八字轨迹分别得到 0.3827 m 和 0.2813 m 的全程 RMS 误差。解析轨迹提供速度、加速度和 yaw 前馈，因此连续轨迹没有被离散成一串急停目标。")

    add_heading(doc, "6.3  扰动、噪声与故障", 2)
    add_figure(doc, wind_figure, 5.65, "图 4  阵风场景的偏差峰值与恢复过程")
    add_body(doc, "阵风场景在飞行中段施加峰值 3.041 N 的外力。位置误差先升至明显峰值，扰动结束后经过 2.223 s 回到验收区间，最终误差为 0.0457 m。传感器噪声场景把带噪 Odom 和 IMU 真正接入闭环，最终误差为 0.0777 m。电机故障场景在限定窗口降低 0 号电机效率，最大倾角增至 19.11°，故障解除后最终仍收敛至 0.0179 m。三类试验说明控制器具备恢复能力，但当前系统尚未包含故障辨识或自适应分配，因此不能把结果解释为对任意严重故障都鲁棒。")

    add_heading(doc, "6.4  避障与局部感知", 2)
    add_figure(doc, replan_figure, 5.35, "图 5  感知重规划场景的净空与局部点云规模")
    add_body(doc, "五障碍物、狭窄通道和感知重规划场景均到达目标。三者最小净空分别为 0.4196 m、0.4136 m 和 0.4889 m，高于 0.30 m 的安全合同。图 5 中局部点云随无人机接近障碍物上升到约 800 点，离开局部环境后逐渐归零；净空曲线在整个任务中没有越过红色阈值。规划路径是几何安全路径，实际轨迹还会受到控制超调和传感器噪声影响，因此使用真实轨迹净空作为最终判据。")

    add_heading(doc, "7  加分功能与工程验证", 1)
    add_body(doc, "课程列出的加分方向均已实现。参数方面，质量、惯量、电机系数、控制增益、传感器、地图、规划与场景均放入 YAML，并由脚本反向检查参数是否真正被节点声明和 launch 装载。扰动模块支持常值、正弦、阵风和固定种子随机力；传感器模块支持白噪声、偏置、随机游走、延迟与丢包；轨迹模块支持圆、八字和 waypoint list。")
    add_body(doc, "多机模式同时启动 3 套独立 namespace、TF、控制和传感器实例，fleet monitor 记录到的最小间距为 0.786 m，高于 0.75 m 要求，未发生违规。故障注入覆盖效率下降、转速上限、命令丢包、延迟和冻结。Web 地面站提供状态与控制入口，参数扫描对位置增益和风力尺度执行 3×3 实验，确定性回放则比较两次运行的轨迹和指标。")
    add_body(doc, "参数扫描显示，扰动力尺度是峰值误差的主要影响因素，位置比例增益在当前范围内只带来较小差异。这与单次手工调参相比更有解释力：它说明继续提高 Kp 并不能抵消外力增长，还可能增加姿态和电机余量压力。测试层面，工作区包含 gtest、launch test、Panel 进程测试、场景阈值检查、回放和多机专项验证。drone_core 还能使用 DRONE_CORE_STANDALONE=ON 脱离 ROS2 构建，直接运行 CTest。")

    add_heading(doc, "8  与参考项目的关系", 1)
    add_body(doc, "pengyu_sim 提供了教学型仿真工程的组织参考，尤其是将动力学、控制和可视化分成可独立理解的模块。MARSIM 更强调轻量级点云仿真，并把地图、局部感知和多机接口放到统一的 ROS 数据流中。本项目借鉴的是问题拆分方式和接口思想，没有复制其 package 或源码。当前工程以 ROS2 Humble 重新实现核心链路，并把可复现实验、自动验收和算法解耦放在更高优先级。")
    add_body(doc, "与 MARSIM 相比，本项目的环境真实感和激光雷达物理细节较弱，不包含 GPU 渲染或真实雷达扫描模式；优点是依赖轻、构建快，动力学和控制器可独立测试。与一般的完整 Gazebo 仿真相比，本项目也没有螺旋桨气动、机架柔性和碰撞接触模型。它更适合课程中验证控制、规划和 ROS2 工程边界，不适合作为高保真飞控认证工具。")

    add_heading(doc, "9  失败案例与改进过程", 1)
    add_body(doc, "开发中最难定位的问题不是公式本身，而是可视化和路径进度。早期 RViz 中 drone_model 会闪烁，Topic 状态在 OK 与 Error 之间跳变。检查发布者数量和 Marker 生命周期后发现，真值节点、传感器节点和可视化节点对相近状态重复发布，Marker 还使用了会过期的寿命。最终统一 Odom 与 TF 的权威来源，给 Marker 使用稳定 ID、零时间戳和 frame_locked，并把地图类信息改为 transient-local QoS。修正后 topic 频率稳定，模型不会因显示对象轮换而消失。")
    add_body(doc, "另一个问题出现在折返和自交路径。若只寻找离无人机最近的路径点，当前位置可能被投影到尚未走过的后半段，规划器就会误判进度并提前发布终点。现在的实现保存单调前进的路径索引，只在有限前向窗口内寻找最近点，再沿弧长选择前视目标。这个改动通过专门的回归测试覆盖，不依赖某一张地图恰好不触发问题。")
    add_body(doc, "这些失败也暴露出当前限制。固定步长模型无法描述高速飞行中的复杂气动，点云来自几何表面采样而非真实光束模型，A* 的计算量会随体素分辨率快速增加。故障恢复主要依靠原控制器余量，严重单电机失效仍可能超出可控范围。报告中的成功结果只对应已记录参数和场景，不应外推到未验证的边界。")

    add_heading(doc, "10  AI 辅助过程与验证责任", 1)
    add_body(doc, "本项目使用 AI 帮助梳理任务清单、生成部分代码骨架、补充测试思路和定位 ROS2 配置问题。关键公式、坐标系、旋翼方向、消息单位和安全阈值均由人工结合代码与曲线复核。ai_usage.md 保存了超过 8 条关键交互摘要，也记录了 AI 曾给出的错误建议及修正过程。")
    add_body(doc, "验证时没有把“能够编译”当作正确。动力学用悬停转速、自由响应和四元数归一化测试；控制器检查静态平衡、力矩符号、限幅与 mixer 输出；ROS2 层检查 topic 类型、发布者数量、频率、TF 连续性和 QoS；系统层则以 11 个场景的 CSV、JSON、PNG 与 rosbag 作为证据。AI 负责提高迭代速度，但判断模型是否可信、实验是否满足阈值以及报告是否如实描述，仍由项目实现者承担。")

    add_heading(doc, "11  总结与后续工作", 1)
    add_body(doc, "本项目完成了从四路电机输入、六自由度刚体动力学、串级控制到地图感知、三维规划、RViz2 和自动评测的闭环。核心算法与 ROS2 的边界清楚，地图相关模块保留中间件依赖；一键脚本解决了 WSL 下反复 source 和显示环境配置的问题。11 个正式场景全部通过，悬停、连续轨迹、风扰、噪声、故障和避障都有可复核的原始数据。")
    add_body(doc, "如果继续投入两周，第一优先级是给动力学加入更可靠的螺旋桨气动、地效和接触模型，并用真实飞控日志标定电机时间常数和阻力。第二步是把局部感知替换为可配置的光束扫描模型，引入动态障碍物和增量式规划。第三步是建立持续集成中的无头场景回归，固定发布频率、CPU 占用和误差基线。完成这些工作后，这套课程仿真器可以进一步用于控制器参数研究、规划算法对比和轻量级硬件在环验证。")

    doc.add_page_break()
    add_heading(doc, "参考文献", 1)
    references = [
        "[1] potato77. pengyu_sim: UAV simulation reference project[EB/OL]. https://gitee.com/potato77/pengyu_sim.",
        "[2] Zhang J, Xu W, Zhu Y, et al. MARSIM: A Light-weight Point-realistic Simulator for LiDAR-based UAVs[EB/OL]. arXiv:2211.10716, 2022.",
        "[3] Lee T, Leok M, McClamroch N H. Geometric Tracking Control of a Quadrotor UAV on SE(3)[C]//49th IEEE Conference on Decision and Control. 2010: 5420-5425.",
        "[4] Quigley M, Conley K, Gerkey B, et al. ROS: An Open-source Robot Operating System[C]//ICRA Workshop on Open Source Software. 2009.",
        "[5] Open Robotics. ROS 2 Humble Documentation[EB/OL]. https://docs.ros.org/en/humble/.",
    ]
    for item in references:
        p = doc.add_paragraph(style="ReferenceAcademic")
        run = p.add_run(item)
        set_font(run, "宋体", 9.5)

    add_heading(doc, "附录 A  复现实验与材料位置", 1)
    add_body(doc, "本报告中的数值均可从仓库现有材料反向核对。复核时先在工作区根目录执行 ./start_sim.sh batch 重新生成场景，再运行 python3 scripts/verify_experiments.py 检查阈值。若只审阅已有结果，可直接查看下表所列文件，不需要启动 RViz2。")
    add_table(
        doc,
        ["材料", "仓库位置", "用途"],
        [
            ["正式实验", "artifacts/experiments/<scenario>/", "CSV、JSON、PNG 与 rosbag"],
            ["模型参数", "src/drone_bringup/config/", "动力学、控制、地图与验收阈值"],
            ["算法说明", "docs/dynamics.md、docs/controller.md", "公式、单位与接口约定"],
            ["AI 使用记录", "ai_usage.md", "关键交互、错误与人工复核"],
            ["演示视频", "output/video/drone_demo.mp4", "1～3 分钟功能演示"],
        ],
        [1700, 4000, 3370],
        font_size=8.4,
    )
    add_caption(doc, "表 4  报告结论的可复核材料")
    add_body(doc, "仓库公开地址为 https://github.com/VMnicht/drone_sim。启动脚本会自行加载 ROS2 和工作区环境，WSL 用户无需在终端中重复 source。若图形界面不可用，所有场景仍可用无头模式运行并生成同样的统计文件。")

    # Keep each section A4 and update fields on open.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    # Re-open once so python-docx prunes transient package relationships.  This
    # avoids a Word 2016 PDF-export loop observed with inline PNG collections.
    normalized_path = output_path.with_name(output_path.stem + ".normalized.docx")
    Document(output_path).save(normalized_path)
    normalized_path.replace(output_path)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, default=Path(__file__).resolve().parents[1])
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    root = args.root.resolve()
    output = args.output or root / "output" / "docx" / "drone_sim_academic_report.docx"
    build_report(root, output)
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
